import csv
import io
import re
from dataclasses import dataclass, field
from pathlib import Path


@dataclass(slots=True)
class ParsedDocument:
    text: str
    metadata: dict = field(default_factory=dict)


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_document_content(filename: str, content: bytes) -> ParsedDocument:
    suffix = Path(filename).suffix.lower().lstrip(".")
    metadata = {"filename": filename, "file_type": suffix}

    if suffix == "pdf":
        return _parse_pdf(content, metadata)
    if suffix == "docx":
        return _parse_docx(content, metadata)
    if suffix == "csv":
        return _parse_csv(content, metadata)
    if suffix in {"md", "markdown", "txt"}:
        return ParsedDocument(clean_text(content.decode("utf-8", errors="ignore")), metadata)
    return ParsedDocument(clean_text(content.decode("utf-8", errors="ignore")), metadata)


def parse_document_file(path: str) -> ParsedDocument:
    file_path = Path(path)
    return parse_document_content(file_path.name, file_path.read_bytes())


def _parse_pdf(content: bytes, metadata: dict) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is installed in runtime image
        raise RuntimeError("pypdf is required to parse PDF files") from exc

    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {page_number}]\n{text}")
    metadata["page_count"] = len(reader.pages)
    return ParsedDocument(clean_text("\n\n".join(pages)), metadata)


def _parse_docx(content: bytes, metadata: dict) -> ParsedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover - dependency is installed in runtime image
        raise RuntimeError("python-docx is required to parse DOCX files") from exc

    document = DocxDocument(io.BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    metadata["paragraph_count"] = len(paragraphs)
    return ParsedDocument(clean_text("\n\n".join(paragraphs)), metadata)


def _parse_csv(content: bytes, metadata: dict) -> ParsedDocument:
    decoded = content.decode("utf-8-sig", errors="ignore")
    reader = csv.DictReader(io.StringIO(decoded))
    rows = []
    for index, row in enumerate(reader, start=1):
        values = [f"{key}: {value}" for key, value in row.items()]
        rows.append(f"Row {index}. " + "; ".join(values))
    metadata["row_count"] = len(rows)
    if rows:
        return ParsedDocument(clean_text("\n".join(rows)), metadata)
    return ParsedDocument(clean_text(decoded), metadata)
